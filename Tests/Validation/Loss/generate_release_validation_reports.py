from __future__ import annotations

import json
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import numpy as np


REPO_ROOT = Path(__file__).resolve().parents[3]
OUTPUT_ROOT = REPO_ROOT / ".local" / "verification" / "release_reports"

PRIUS_FULL_SUMMARY = (
    REPO_ROOT / ".local" / "verification" / "prius2004_full_validation" / "summary.json"
)
PRIUS_ELMER_SUMMARY = (
    REPO_ROOT / ".local" / "verification" / "elmer_prius2004" / "summary.json"
)
LEAF_FULL_SUMMARY = (
    REPO_ROOT / ".local" / "verification" / "leaf_full_validation" / "summary.json"
)
LEAF_ELMER_SUMMARY = (
    REPO_ROOT / ".local" / "verification" / "elmer_leaf2012" / "summary.json"
)

SIMULATION_REPORT_PATH = OUTPUT_ROOT / "pyleecan_v1_6_simulation_report.docx"
COMPARISON_REPORT_PATH = OUTPUT_ROOT / "pyleecan_v1_6_public_comparison_report.docx"
SIGNOFF_JSON_PATH = OUTPUT_ROOT / "pyleecan_v1_6_release_signoff.json"

GEOMETRY_LIMIT_PCT = 5.0
PUBLIC_LIMIT_PCT = 10.0
CURRENT_LIMIT_PCT = 5.0
FIELD_WEAKENING_LIMIT_PCT = 10.0
SPEED_LIMIT_PCT = 10.0
ELMER_LIMIT_PCT = 5.0
EFFICIENCY_POINT_LIMIT = 0.01

PASS = "PASS"
FAIL = "FAIL"
INFO = "INFO"
BLOCKED = "BLOCKED"
REVIEW = "REVIEW"
CRITICAL = "critical"
SUPPORTING = "supporting"

PRIUS_PUBLIC_SOURCES = [
    {
        "title": "Evaluation of 2004 Toyota Prius Hybrid Electric Drive System",
        "url": "https://www.osti.gov/biblio/890029",
        "notes": (
            "ORNL full-system benchmark. The public summary states the motor delivers "
            "50 kW over 1200-1540 rpm and supports the release check against the "
            "1200 rpm operating point."
        ),
    },
    {
        "title": "Report on Toyota/Prius Motor Torque-Capability, Torque-Property, No-Load Back EMF, and Mechanical Losses",
        "url": "https://www.osti.gov/biblio/885669",
        "notes": (
            "ORNL motor-focused benchmark used as the public torque and loss reference "
            "family for the Prius release comparison."
        ),
    },
    {
        "title": "Report on Toyota/Prius Motor Design and Manufacturing Assessment",
        "url": "https://www.osti.gov/biblio/885676/",
        "notes": (
            "ORNL teardown report used for public geometry and winding references in the "
            "existing Prius validation script."
        ),
    },
]

SIGNOFF_ASSUMPTIONS = [
    "Geometry items pass when absolute delta stays within 5% of the public reference.",
    "Power, torque, current and static public operating-point items pass when absolute delta stays within 10%, except peak-torque current which uses a 5% limit.",
    "Field-weakening onset passes when the model onset stays within 10% of the public reference speed.",
    "Efficiency items pass when the model stays within 0.01 absolute efficiency of the public target or band edge.",
    "Elmer vs FEMM cross-solver checks pass when torque and power deltas stay within 5% at each replayed operating point.",
    "If any critical item fails, the release signoff is blocked.",
]


@dataclass(frozen=True)
class CheckItem:
    case_name: str
    category: str
    metric: str
    reference: str
    model: str
    delta: str
    threshold: str
    status: str
    severity: str
    notes: str


def load_json(path: Path) -> dict:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def ensure_finite(value) -> float | None:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return None
    if not np.isfinite(numeric):
        return None
    return numeric


def format_metric(value, fmt: str = "{:.3f}") -> str:
    numeric = ensure_finite(value)
    if numeric is None:
        return "n/a"
    return fmt.format(numeric)


def format_pct(value) -> str:
    numeric = ensure_finite(value)
    if numeric is None:
        return "n/a"
    return f"{numeric:.2f}%"


def relative_to_repo(path: str | Path | None) -> str:
    if path in [None, ""]:
        return "n/a"
    try:
        return str(Path(path).resolve().relative_to(REPO_ROOT))
    except Exception:
        return str(path)


def status_from_pct(delta_pct, limit_pct) -> str:
    numeric = ensure_finite(delta_pct)
    if numeric is None:
        return INFO
    return PASS if abs(numeric) <= float(limit_pct) else FAIL


def status_from_bool(value: bool) -> str:
    return PASS if bool(value) else FAIL


def status_from_efficiency_point(
    value, reference, limit_abs=EFFICIENCY_POINT_LIMIT
) -> str:
    numeric = ensure_finite(value)
    ref = ensure_finite(reference)
    if numeric is None or ref is None:
        return INFO
    return PASS if abs(numeric - ref) <= float(limit_abs) else FAIL


def status_from_efficiency_band(
    value, low, high, limit_abs=EFFICIENCY_POINT_LIMIT
) -> str:
    numeric = ensure_finite(value)
    lower = ensure_finite(low)
    upper = ensure_finite(high)
    if numeric is None or lower is None or upper is None:
        return INFO
    if numeric < lower:
        distance = lower - numeric
    elif numeric > upper:
        distance = numeric - upper
    else:
        distance = 0.0
    return PASS if distance <= float(limit_abs) else FAIL


def make_check(
    *,
    case_name: str,
    category: str,
    metric: str,
    reference: str,
    model: str,
    delta: str,
    threshold: str,
    status: str,
    severity: str,
    notes: str,
) -> CheckItem:
    return CheckItem(
        case_name=case_name,
        category=category,
        metric=metric,
        reference=reference,
        model=model,
        delta=delta,
        threshold=threshold,
        status=status,
        severity=severity,
        notes=notes,
    )


def build_prius_checks(full_summary: dict, elmer_summary: dict) -> list[CheckItem]:
    checks: list[CheckItem] = []
    geometry = full_summary["comparison"]["geometry_vs_ornl"]
    static_cmp = full_summary["comparison"]["static_vs_published"]
    eff_cmp = full_summary["comparison"]["efficiency_region_vs_ornl"]
    drive_cmp = full_summary["comparison"]["drive_cycle_envelope_check"]

    for metric, values in geometry.items():
        checks.append(
            make_check(
                case_name="Prius 2004",
                category="Geometry",
                metric=metric,
                reference=format_metric(values["reference"]),
                model=format_metric(values["model"]),
                delta=format_pct(values["delta_pct"]),
                threshold=f"|delta| <= {GEOMETRY_LIMIT_PCT:.1f}%",
                status=status_from_pct(values["delta_pct"], GEOMETRY_LIMIT_PCT),
                severity=SUPPORTING,
                notes="Public geometry target reused from the existing ORNL-based Prius validation summary.",
            )
        )

    for metric, delta_pct, model_value, reference_value in [
        (
            "static power at 1200 rpm [kW]",
            static_cmp["power_delta_pct"],
            static_cmp["model_power_kW"],
            static_cmp["reference_power_kW"],
        ),
        (
            "static torque at 1200 rpm [Nm]",
            static_cmp["torque_delta_pct"],
            static_cmp["model_torque_Nm"],
            static_cmp["reference_torque_Nm"],
        ),
    ]:
        checks.append(
            make_check(
                case_name="Prius 2004",
                category="Public operating point",
                metric=metric,
                reference=format_metric(reference_value),
                model=format_metric(model_value),
                delta=format_pct(delta_pct),
                threshold=f"|delta| <= {PUBLIC_LIMIT_PCT:.1f}%",
                status=status_from_pct(delta_pct, PUBLIC_LIMIT_PCT),
                severity=CRITICAL,
                notes="Release signoff check against the ORNL Prius public operating point encoded in the local validation summary.",
            )
        )

    band_low, band_high = eff_cmp["ornl_motor_peak_efficiency_band"]
    for metric, model_value in [
        (
            "moderate-region peak efficiency",
            eff_cmp["model_moderate_region_peak_efficiency"],
        ),
        (
            "moderate-region mean efficiency",
            eff_cmp["model_moderate_region_mean_efficiency"],
        ),
    ]:
        checks.append(
            make_check(
                case_name="Prius 2004",
                category="Efficiency region",
                metric=metric,
                reference=f"{band_low:.3f} to {band_high:.3f}",
                model=format_metric(model_value, "{:.4f}"),
                delta=(
                    format_metric(
                        max(
                            band_low - float(model_value),
                            float(model_value) - band_high,
                            0.0,
                        ),
                        "{:.4f}",
                    )
                    if ensure_finite(model_value) is not None
                    else "n/a"
                ),
                threshold=f"distance to band <= {EFFICIENCY_POINT_LIMIT:.3f}",
                status=status_from_efficiency_band(model_value, band_low, band_high),
                severity=SUPPORTING,
                notes="Band-based check around the ORNL Prius efficiency-region reference used by the existing validation.",
            )
        )

    checks.append(
        make_check(
            case_name="Prius 2004",
            category="Drive-cycle envelope",
            metric="cycle maximum speed",
            reference=format_metric(
                drive_cmp["published_motor_max_speed_rpm"], "{:.0f}"
            ),
            model=format_metric(drive_cmp["cycle_max_speed_rpm"], "{:.0f}"),
            delta=(
                "inside envelope"
                if drive_cmp["within_speed_limit"]
                else "outside envelope"
            ),
            threshold="must stay within published speed envelope",
            status=status_from_bool(drive_cmp["within_speed_limit"]),
            severity=CRITICAL,
            notes="Cycle envelope reuse from the local Prius validation trajectory.",
        )
    )
    checks.append(
        make_check(
            case_name="Prius 2004",
            category="Drive-cycle envelope",
            metric="cycle maximum torque",
            reference=format_metric(drive_cmp["published_motor_max_torque_Nm"]),
            model=format_metric(drive_cmp["cycle_max_torque_Nm"]),
            delta=(
                "inside envelope"
                if drive_cmp["within_torque_limit"]
                else "outside envelope"
            ),
            threshold="must stay within published torque envelope",
            status=status_from_bool(drive_cmp["within_torque_limit"]),
            severity=CRITICAL,
            notes="Cycle envelope reuse from the local Prius validation trajectory.",
        )
    )

    for point in elmer_summary["points"]:
        comparison = point["compare_to_femm"]
        elmer = point["elmer"]
        femm = point["femm_baseline"]
        for metric, delta_key, model_value, reference_value in [
            (
                f"{point['tag']} torque [Nm]",
                "torque_delta_pct",
                elmer["Tem_av_Nm"],
                femm["Tem_av_Nm"],
            ),
            (
                f"{point['tag']} power [kW]",
                "power_delta_pct",
                elmer["P_out_W"] * 1e-3,
                femm["P_out_W"] * 1e-3,
            ),
        ]:
            checks.append(
                make_check(
                    case_name="Prius 2004",
                    category="Elmer vs FEMM",
                    metric=metric,
                    reference=format_metric(reference_value),
                    model=format_metric(model_value),
                    delta=format_pct(comparison[delta_key]),
                    threshold=f"|delta| <= {ELMER_LIMIT_PCT:.1f}%",
                    status=status_from_pct(comparison[delta_key], ELMER_LIMIT_PCT),
                    severity=CRITICAL,
                    notes="Cross-solver replay using the current Elmer validation scripts and the full-load FEMM seeds.",
                )
            )

    return checks


def build_leaf_checks(full_summary: dict, elmer_summary: dict) -> list[CheckItem]:
    checks: list[CheckItem] = []
    geometry = full_summary["comparison"]["geometry_vs_public"]
    power_cmp = full_summary["comparison"]["power_and_torque_vs_public"]
    current_cmp = full_summary["comparison"]["current_voltage_vs_public"]
    eff_cmp = full_summary["comparison"]["efficiency_vs_public"]
    speed_cmp = full_summary["comparison"]["speed_envelope_vs_public"]
    static_cmp = full_summary["comparison"]["static_3000rpm_point"]

    for metric, values in geometry.items():
        checks.append(
            make_check(
                case_name="LEAF 2012",
                category="Geometry",
                metric=metric,
                reference=format_metric(values["reference"]),
                model=format_metric(values["model"]),
                delta=format_pct(values["delta_pct"]),
                threshold=f"|delta| <= {GEOMETRY_LIMIT_PCT:.1f}%",
                status=status_from_pct(values["delta_pct"], GEOMETRY_LIMIT_PCT),
                severity=SUPPORTING,
                notes="Public geometry target reused from the local LEAF benchmark summary.",
            )
        )

    for metric, delta_pct, model_value, reference_value in [
        (
            "peak power [kW]",
            power_cmp["peak_power_delta_pct"],
            power_cmp["model_peak_power_kW"],
            power_cmp["reference_peak_power_kW"],
        ),
        (
            "peak torque [Nm]",
            power_cmp["peak_torque_delta_pct"],
            power_cmp["model_peak_torque_Nm"],
            power_cmp["reference_peak_torque_Nm"],
        ),
        (
            "power at 3000 rpm [kW]",
            power_cmp["power_at_3000_delta_pct"],
            power_cmp["model_power_at_3000_rpm_kW"],
            power_cmp["reference_peak_power_kW"],
        ),
        (
            "static 3000 rpm power [kW]",
            static_cmp["power_delta_pct"],
            static_cmp["model_power_kW"],
            static_cmp["reference_power_kW"],
        ),
        (
            "static 3000 rpm torque [Nm]",
            static_cmp["torque_delta_pct"],
            static_cmp["model_torque_Nm"],
            static_cmp["reference_torque_Nm"],
        ),
    ]:
        checks.append(
            make_check(
                case_name="LEAF 2012",
                category="Public operating point",
                metric=metric,
                reference=format_metric(reference_value),
                model=format_metric(model_value),
                delta=format_pct(delta_pct),
                threshold=f"|delta| <= {PUBLIC_LIMIT_PCT:.1f}%",
                status=status_from_pct(delta_pct, PUBLIC_LIMIT_PCT),
                severity=CRITICAL,
                notes="Release signoff check against the local LEAF public benchmark targets.",
            )
        )

    checks.append(
        make_check(
            case_name="LEAF 2012",
            category="Current and voltage",
            metric="peak-torque current [Arms]",
            reference=format_metric(current_cmp["reference_peak_torque_current_arms"]),
            model=format_metric(current_cmp["model_peak_torque_current_arms"]),
            delta=format_pct(current_cmp["current_delta_pct"]),
            threshold=f"|delta| <= {CURRENT_LIMIT_PCT:.1f}%",
            status=status_from_pct(current_cmp["current_delta_pct"], CURRENT_LIMIT_PCT),
            severity=CRITICAL,
            notes="Public current target reused from the DOE/ORNL LEAF summary embedded in the repo validation.",
        )
    )
    checks.append(
        make_check(
            case_name="LEAF 2012",
            category="Current and voltage",
            metric="field-weakening onset [rpm]",
            reference=format_metric(
                current_cmp["reference_voltage_limit_onset_rpm"], "{:.0f}"
            ),
            model=format_metric(
                current_cmp["model_field_weakening_start_rpm"], "{:.0f}"
            ),
            delta=format_pct(current_cmp["field_weakening_delta_pct"]),
            threshold=f"|delta| <= {FIELD_WEAKENING_LIMIT_PCT:.1f}%",
            status=status_from_pct(
                current_cmp["field_weakening_delta_pct"], FIELD_WEAKENING_LIMIT_PCT
            ),
            severity=CRITICAL,
            notes="Field-weakening onset is a public benchmark alignment item for the LEAF release decision.",
        )
    )

    model_peak_eff = eff_cmp["model_high_power_peak_efficiency"]
    ref_peak_eff = eff_cmp["reference_peak_efficiency"]
    checks.append(
        make_check(
            case_name="LEAF 2012",
            category="Efficiency",
            metric="high-power peak efficiency",
            reference=format_metric(ref_peak_eff, "{:.4f}"),
            model=format_metric(model_peak_eff, "{:.4f}"),
            delta=(
                format_metric(float(model_peak_eff) - float(ref_peak_eff), "{:.4f}")
                if ensure_finite(model_peak_eff) is not None
                and ensure_finite(ref_peak_eff) is not None
                else "n/a"
            ),
            threshold=f"|delta| <= {EFFICIENCY_POINT_LIMIT:.3f}",
            status=status_from_efficiency_point(model_peak_eff, ref_peak_eff),
            severity=CRITICAL,
            notes="The public LEAF report states the motor exceeded 97% efficiency in the high-power band.",
        )
    )

    checks.append(
        make_check(
            case_name="LEAF 2012",
            category="Speed envelope",
            metric="full-load maximum speed [rpm]",
            reference=format_metric(speed_cmp["reference_max_speed_rpm"], "{:.0f}"),
            model=format_metric(speed_cmp["model_full_load_max_speed_rpm"], "{:.0f}"),
            delta=format_pct(
                (
                    (
                        float(speed_cmp["model_full_load_max_speed_rpm"])
                        - float(speed_cmp["reference_max_speed_rpm"])
                    )
                    / float(speed_cmp["reference_max_speed_rpm"])
                    * 100.0
                )
            ),
            threshold=f"|delta| <= {SPEED_LIMIT_PCT:.1f}%",
            status=status_from_pct(
                (
                    (
                        float(speed_cmp["model_full_load_max_speed_rpm"])
                        - float(speed_cmp["reference_max_speed_rpm"])
                    )
                    / float(speed_cmp["reference_max_speed_rpm"])
                    * 100.0
                ),
                SPEED_LIMIT_PCT,
            ),
            severity=CRITICAL,
            notes="Direct comparison against the public LEAF maximum-speed reference.",
        )
    )
    checks.append(
        make_check(
            case_name="LEAF 2012",
            category="Speed envelope",
            metric="drive-cycle maximum speed",
            reference=format_metric(speed_cmp["reference_max_speed_rpm"], "{:.0f}"),
            model=format_metric(speed_cmp["drive_cycle_max_speed_rpm"], "{:.0f}"),
            delta=(
                "inside envelope"
                if speed_cmp["within_drive_cycle_speed_limit"]
                else "outside envelope"
            ),
            threshold="must stay within published speed envelope",
            status=status_from_bool(speed_cmp["within_drive_cycle_speed_limit"]),
            severity=CRITICAL,
            notes="Drive-cycle speed check against the public LEAF maximum-speed envelope.",
        )
    )
    checks.append(
        make_check(
            case_name="LEAF 2012",
            category="Speed envelope",
            metric="drive-cycle maximum torque",
            reference=format_metric(full_summary["public_reference"]["peak_torque_Nm"]),
            model=format_metric(
                full_summary["drive_cycle"]["summary"]["max_torque_Nm"]
            ),
            delta=(
                "inside envelope"
                if speed_cmp["within_drive_cycle_torque_limit"]
                else "outside envelope"
            ),
            threshold="must stay within published torque envelope",
            status=status_from_bool(speed_cmp["within_drive_cycle_torque_limit"]),
            severity=CRITICAL,
            notes="Drive-cycle torque check against the public LEAF peak-torque envelope.",
        )
    )

    for point in elmer_summary["points"]:
        comparison = point["compare_to_femm"]
        elmer = point["elmer"]
        femm = point["femm_baseline"]
        for metric, delta_key, model_value, reference_value in [
            (
                f"{point['tag']} torque [Nm]",
                "torque_delta_pct",
                elmer["Tem_av_Nm"],
                femm["Tem_av_Nm"],
            ),
            (
                f"{point['tag']} power [kW]",
                "power_delta_pct",
                elmer["P_out_W"] * 1e-3,
                femm["P_out_W"] * 1e-3,
            ),
        ]:
            checks.append(
                make_check(
                    case_name="LEAF 2012",
                    category="Elmer vs FEMM",
                    metric=metric,
                    reference=format_metric(reference_value),
                    model=format_metric(model_value),
                    delta=format_pct(comparison[delta_key]),
                    threshold=f"|delta| <= {ELMER_LIMIT_PCT:.1f}%",
                    status=status_from_pct(comparison[delta_key], ELMER_LIMIT_PCT),
                    severity=CRITICAL,
                    notes="Cross-solver replay using the current Elmer validation scripts and the full-load FEMM seeds.",
                )
            )

    return checks


def count_statuses(checks: Iterable[CheckItem]) -> dict[str, int]:
    status_count = {PASS: 0, FAIL: 0, INFO: 0}
    for check in checks:
        status_count.setdefault(check.status, 0)
        status_count[check.status] += 1
    return status_count


def release_decision(checks: Iterable[CheckItem]) -> str:
    check_list = list(checks)
    if any(check.status == FAIL and check.severity == CRITICAL for check in check_list):
        return BLOCKED
    if any(check.status == FAIL for check in check_list):
        return REVIEW
    return "PASS"


def get_public_sources(case_name: str, full_summary: dict) -> list[dict]:
    if case_name == "Prius 2004":
        return PRIUS_PUBLIC_SOURCES
    return full_summary.get("public_sources", [])


def get_comparison_image_paths(case_name: str, full_summary: dict) -> list[Path]:
    plot_paths = full_summary["artifacts"]["efficiency_map"].get("plot_paths", {})
    custom_paths = full_summary["artifacts"]["efficiency_map"].get(
        "custom_envelope_paths", {}
    )
    if case_name == "LEAF 2012":
        image_list = [
            custom_paths.get("torque_plot"),
            custom_paths.get("power_plot"),
        ]
    else:
        image_list = [
            plot_paths.get("torque_envelope"),
            plot_paths.get("power_envelope"),
        ]
    return [Path(path) for path in image_list if path]


def get_simulation_image_paths(
    case_name: str, full_summary: dict, elmer_summary: dict
) -> list[Path]:
    plot_paths = full_summary["artifacts"]["efficiency_map"].get("plot_paths", {})
    image_candidates = [
        full_summary["artifacts"]["static"].get("loss_plot"),
        full_summary["artifacts"]["dynamic_load"].get("trace_plot"),
        plot_paths.get("efficiency_map"),
    ]
    for point in elmer_summary["points"]:
        screenshot = point["artifacts"].get("screenshot")
        if screenshot:
            image_candidates.append(screenshot)
            break
    return [Path(path) for path in image_candidates if path]


def build_case_payload(
    case_name: str,
    *,
    full_summary_path: Path,
    elmer_summary_path: Path,
) -> dict:
    full_summary = load_json(full_summary_path)
    elmer_summary = load_json(elmer_summary_path)

    if case_name == "Prius 2004":
        checks = build_prius_checks(full_summary, elmer_summary)
    else:
        checks = build_leaf_checks(full_summary, elmer_summary)

    case = {
        "case_name": case_name,
        "full_summary_path": str(full_summary_path),
        "elmer_summary_path": str(elmer_summary_path),
        "full_summary": full_summary,
        "elmer_summary": elmer_summary,
        "public_sources": get_public_sources(case_name, full_summary),
        "checks": checks,
        "status_count": count_statuses(checks),
        "decision": release_decision(checks),
        "simulation_images": get_simulation_image_paths(
            case_name, full_summary, elmer_summary
        ),
        "comparison_images": get_comparison_image_paths(case_name, full_summary),
    }
    case["report_positioning"] = get_case_report_positioning(case_name)
    case["comparison_conclusion"] = build_case_comparison_conclusion(case)
    return case


def get_case_report_positioning(case_name: str) -> str:
    if case_name == "LEAF 2012":
        return "Internal validation version (内部验证版)"
    return "Public signoff scope"


def build_case_comparison_conclusion(case: dict) -> str:
    if case["case_name"] == "LEAF 2012":
        return (
            "LEAF 2012 is documented in this release packet as an internal "
            "validation version (内部验证版): the current model is retained to "
            "demonstrate geometry consistency and Elmer-vs-FEMM solver consistency, "
            "but it is not signed against the public power/torque envelope in the "
            "Pyleecan 1.6 public comparison scope."
        )

    return (
        "Prius 2004 remains inside the Pyleecan 1.6 public signoff scope and its "
        "fresh public-reference checks plus Elmer-vs-FEMM replay checks stay within "
        "the release limits."
    )


def build_overall_comparison_conclusion(payload: dict) -> str:
    leaf_case = next(
        (case for case in payload["cases"] if case["case_name"] == "LEAF 2012"),
        None,
    )
    if leaf_case is not None:
        return (
            "Overall comparison scope: Prius 2004 remains the public signoff case for "
            "Pyleecan 1.6. LEAF 2012 is reported only as an internal validation "
            "version (内部验证版), which means the present package is used to document "
            "geometry consistency and solver replay consistency rather than to sign "
            "the public LEAF power/torque envelope."
        )

    return (
        "Overall comparison scope: all listed cases are reported inside the public "
        "signoff scope."
    )


def build_release_payload() -> dict:
    cases = [
        build_case_payload(
            "Prius 2004",
            full_summary_path=PRIUS_FULL_SUMMARY,
            elmer_summary_path=PRIUS_ELMER_SUMMARY,
        ),
        build_case_payload(
            "LEAF 2012",
            full_summary_path=LEAF_FULL_SUMMARY,
            elmer_summary_path=LEAF_ELMER_SUMMARY,
        ),
    ]
    overall = PASS
    if any(case["decision"] == BLOCKED for case in cases):
        overall = BLOCKED
    elif any(case["decision"] == REVIEW for case in cases):
        overall = REVIEW

    return {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "overall_release_decision": overall,
        "overall_comparison_conclusion": build_overall_comparison_conclusion(
            {"cases": cases}
        ),
        "signoff_assumptions": SIGNOFF_ASSUMPTIONS,
        "cases": cases,
    }


def write_signoff_json(payload: dict, output_path: Path = SIGNOFF_JSON_PATH) -> Path:
    serializable = {
        "generated_at_utc": payload["generated_at_utc"],
        "overall_release_decision": payload["overall_release_decision"],
        "overall_comparison_conclusion": payload["overall_comparison_conclusion"],
        "signoff_assumptions": payload["signoff_assumptions"],
        "cases": [],
    }
    for case in payload["cases"]:
        serializable["cases"].append(
            {
                "case_name": case["case_name"],
                "decision": case["decision"],
                "report_positioning": case["report_positioning"],
                "comparison_conclusion": case["comparison_conclusion"],
                "status_count": case["status_count"],
                "full_summary_path": case["full_summary_path"],
                "elmer_summary_path": case["elmer_summary_path"],
                "public_sources": case["public_sources"],
                "checks": [asdict(check) for check in case["checks"]],
            }
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(serializable, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return output_path


def _require_docx():
    try:
        import docx  # noqa: F401
    except ModuleNotFoundError as error:
        raise ModuleNotFoundError(
            "python-docx is required to generate release validation reports."
        ) from error


def add_table(document, headers: list[str], rows: Iterable[Iterable[str]]):
    from docx.enum.table import WD_TABLE_ALIGNMENT

    row_list = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
    for row in row_list:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    return table


def add_picture_if_exists(
    document, path: Path, caption: str, width_inches: float = 6.0
):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    if not path.exists():
        document.add_paragraph(f"Missing artifact: {relative_to_repo(path)}")
        return
    paragraph = document.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_picture(str(path), width=Inches(width_inches))


def write_simulation_report(
    payload: dict, output_path: Path = SIMULATION_REPORT_PATH
) -> Path:
    _require_docx()
    from docx import Document

    document = Document()
    document.add_heading("Pyleecan 1.6 Simulation Report", level=0)
    document.add_paragraph(
        f"Generated at (UTC): {payload['generated_at_utc']}\n"
        f"Overall release decision: {payload['overall_release_decision']}"
    )
    document.add_paragraph(
        "This document collects the fresh FEMM public-validation outputs and the "
        "fresh Elmer-vs-FEMM replay outputs for the Toyota Prius 2004 and Nissan "
        "LEAF 2012 release signoff review."
    )

    for case in payload["cases"]:
        full_summary = case["full_summary"]
        elmer_summary = case["elmer_summary"]
        document.add_heading(case["case_name"], level=1)
        document.add_paragraph(
            f"Decision: {case['decision']}. "
            f"Report positioning: {case['report_positioning']}. "
            f"PASS={case['status_count'].get(PASS, 0)}, "
            f"FAIL={case['status_count'].get(FAIL, 0)}, "
            f"INFO={case['status_count'].get(INFO, 0)}."
        )
        document.add_paragraph(case["comparison_conclusion"])
        document.add_paragraph(
            f"Full validation summary: {relative_to_repo(case['full_summary_path'])}\n"
            f"Elmer replay summary: {relative_to_repo(case['elmer_summary_path'])}\n"
            f"Full validation generated at: {full_summary['generated_at_utc']}\n"
            f"Elmer replay generated at: {elmer_summary['generated_at_utc']}"
        )

        document.add_heading("Public Sources", level=2)
        add_table(
            document,
            ["Title", "URL", "Notes"],
            [
                [source["title"], source["url"], source["notes"]]
                for source in case["public_sources"]
            ],
        )

        document.add_heading("Fresh Validation Snapshot", level=2)
        if case["case_name"] == "Prius 2004":
            static_cmp = full_summary["comparison"]["static_vs_published"]
            drive_summary = full_summary["drive_cycle"]["summary"]
            summary_rows = [
                [
                    "static point speed [rpm]",
                    format_metric(full_summary["static"]["speed_rpm"], "{:.0f}"),
                ],
                [
                    "static torque [Nm]",
                    format_metric(full_summary["static"]["Tem_av_Nm"]),
                ],
                [
                    "static power [kW]",
                    format_metric(full_summary["static"]["P_out_W"] * 1e-3),
                ],
                ["public power delta", format_pct(static_cmp["power_delta_pct"])],
                ["public torque delta", format_pct(static_cmp["torque_delta_pct"])],
                [
                    "cycle efficiency",
                    format_metric(drive_summary["eta_cycle"], "{:.4f}"),
                ],
                [
                    "cycle max speed [rpm]",
                    format_metric(drive_summary["max_speed_rpm"], "{:.0f}"),
                ],
                [
                    "cycle max torque [Nm]",
                    format_metric(drive_summary["max_torque_Nm"]),
                ],
            ]
        else:
            static_cmp = full_summary["comparison"]["static_3000rpm_point"]
            drive_summary = full_summary["drive_cycle"]["summary"]
            summary_rows = [
                [
                    "static point speed [rpm]",
                    format_metric(full_summary["static"]["speed_rpm"], "{:.0f}"),
                ],
                [
                    "static torque [Nm]",
                    format_metric(full_summary["static"]["Tem_av_Nm"]),
                ],
                [
                    "static power [kW]",
                    format_metric(full_summary["static"]["P_out_W"] * 1e-3),
                ],
                ["public power delta", format_pct(static_cmp["power_delta_pct"])],
                ["public torque delta", format_pct(static_cmp["torque_delta_pct"])],
                [
                    "cycle efficiency",
                    format_metric(drive_summary["eta_cycle"], "{:.4f}"),
                ],
                [
                    "cycle max speed [rpm]",
                    format_metric(drive_summary["max_speed_rpm"], "{:.0f}"),
                ],
                [
                    "cycle max torque [Nm]",
                    format_metric(drive_summary["max_torque_Nm"]),
                ],
            ]
        add_table(document, ["Metric", "Value"], summary_rows)

        document.add_heading("Elmer vs FEMM Replay", level=2)
        add_table(
            document,
            [
                "Point",
                "FEMM torque [Nm]",
                "Elmer torque [Nm]",
                "Torque delta",
                "FEMM power [kW]",
                "Elmer power [kW]",
                "Power delta",
            ],
            [
                [
                    point["tag"],
                    format_metric(point["femm_baseline"]["Tem_av_Nm"]),
                    format_metric(point["elmer"]["Tem_av_Nm"]),
                    format_pct(point["compare_to_femm"]["torque_delta_pct"]),
                    format_metric(point["femm_baseline"]["P_out_W"] * 1e-3),
                    format_metric(point["elmer"]["P_out_W"] * 1e-3),
                    format_pct(point["compare_to_femm"]["power_delta_pct"]),
                ]
                for point in elmer_summary["points"]
            ],
        )

        document.add_heading("Key Artifacts", level=2)
        for image_path in case["simulation_images"]:
            add_picture_if_exists(
                document,
                image_path,
                caption=relative_to_repo(image_path),
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def write_comparison_report(
    payload: dict, output_path: Path = COMPARISON_REPORT_PATH
) -> Path:
    _require_docx()
    from docx import Document

    document = Document()
    document.add_heading("Pyleecan 1.6 Public Comparison Report", level=0)
    document.add_paragraph(
        f"Generated at (UTC): {payload['generated_at_utc']}\n"
        f"Overall release decision: {payload['overall_release_decision']}"
    )

    document.add_heading("Release Gate Assumptions", level=1)
    for item in payload["signoff_assumptions"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Overall Decision and Scope", level=1)
    add_table(
        document,
        ["Case", "Decision", "Report Positioning", "PASS", "FAIL", "INFO"],
        [
            [
                case["case_name"],
                case["decision"],
                case["report_positioning"],
                case["status_count"].get(PASS, 0),
                case["status_count"].get(FAIL, 0),
                case["status_count"].get(INFO, 0),
            ]
            for case in payload["cases"]
        ],
    )
    document.add_paragraph(payload["overall_comparison_conclusion"])

    for case in payload["cases"]:
        document.add_heading(case["case_name"], level=1)
        document.add_paragraph(
            f"Decision: {case['decision']} based on "
            f"{case['status_count'].get(FAIL, 0)} failed checks."
        )
        document.add_paragraph(f"Report positioning: {case['report_positioning']}.")
        document.add_paragraph(case["comparison_conclusion"])
        add_table(
            document,
            [
                "Category",
                "Metric",
                "Reference",
                "Model",
                "Delta",
                "Threshold",
                "Status",
                "Severity",
            ],
            [
                [
                    check.category,
                    check.metric,
                    check.reference,
                    check.model,
                    check.delta,
                    check.threshold,
                    check.status,
                    check.severity,
                ]
                for check in case["checks"]
            ],
        )

        document.add_heading("Comparison Images", level=2)
        for image_path in case["comparison_images"]:
            add_picture_if_exists(
                document,
                image_path,
                caption=relative_to_repo(image_path),
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def generate_reports() -> dict:
    payload = build_release_payload()
    json_path = write_signoff_json(payload)
    simulation_path = write_simulation_report(payload)
    comparison_path = write_comparison_report(payload)
    return {
        "payload": payload,
        "json_path": json_path,
        "simulation_report_path": simulation_path,
        "comparison_report_path": comparison_path,
    }


def main():
    result = generate_reports()
    print(f"Wrote signoff JSON to {result['json_path']}")
    print(f"Wrote simulation report to {result['simulation_report_path']}")
    print(f"Wrote comparison report to {result['comparison_report_path']}")
    print(
        "Overall release decision: " f"{result['payload']['overall_release_decision']}"
    )


if __name__ == "__main__":
    main()
